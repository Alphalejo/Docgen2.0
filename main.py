import sys
import os
sys.path.insert(0, os.path.dirname(__file__))

import io
import json
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from weasyprint import HTML
from docx import Document
from fastapi.responses import Response

from templates_config import TEMPLATES

app = FastAPI(title="Document Builder API")

# Mount static files & configure Jinja2 templates
app.mount("/static", StaticFiles(directory="app/static"), name="static")
templates = Jinja2Templates(directory="app/templates")

class PreviewRequest(BaseModel):
    template: str
    data: dict

@app.get("/", response_class=HTMLResponse)
async def serve_ui():
    """Serves the main vanilla HTML interface."""
    with open("app/static/index.html", "r", encoding="utf-8") as f:
        return f.read()

@app.get("/templates")
async def list_templates():
    """Returns available templates config to build the UI form dynamically."""
    return TEMPLATES

@app.post("/preview", response_class=HTMLResponse)
async def preview_document(payload: PreviewRequest):
    """Renders Jinja2 HTML and returns the string for the frontend iframe."""
    if payload.template not in TEMPLATES:
        return HTMLResponse("Invalid Template", status_code=400)
    
    template = templates.get_template(f"{payload.template}.html")
    html_content = template.render(**payload.data)
    return html_content

@app.post("/export/pdf")
async def export_pdf(payload: PreviewRequest):
    template = templates.get_template(f"{payload.template}.html")
    html_content = template.render(**payload.data)
    pdf_bytes = HTML(string=html_content).write_pdf()
    
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={payload.template}.pdf"}
    )

@app.post("/export/docx")
async def export_docx(payload: PreviewRequest):
    """Builds a Word document natively using python-docx."""
    doc = Document()
    data = payload.data
    
    # Custom python-docx mapping per template shape
    if payload.template == "invoice":
        doc.add_heading("INVOICE", 0)
        doc.add_heading(data.get("company_name", "Your Company"), 1)
        doc.add_paragraph(f"Invoice #: {data.get('invoice_number', '')}")
        doc.add_paragraph(f"Date: {data.get('date', '')}")
        doc.add_paragraph(f"Billed To: {data.get('client_name', '')}")
        
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = 'Description'
        table.cell(0, 1).text = 'Total ($)'
        table.cell(1, 0).text = data.get("description", "")
        table.cell(1, 1).text = f"${data.get('amount', '0.00')}"
        
        doc.add_paragraph(f"\nTotal Due: ${data.get('amount', '0.00')}")
        
    elif payload.template == "report":
        doc.add_heading(data.get("title", "Monthly Report"), 0)
        doc.add_paragraph(f"Prepared by: {data.get('author', '')}")
        doc.add_paragraph(f"Date: {data.get('date', '')}")
        
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(data.get("summary", ""))
        
        doc.add_heading("Key Metrics", level=1)
        doc.add_paragraph(data.get("metrics", ""))

    elif payload.template == "cv":
        doc.add_heading(data.get("name", "Full Name"), 0)
        doc.add_paragraph(data.get("title", ""))
        doc.add_paragraph(f"{data.get('email', '')} | {data.get('location', '')}")
        
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(data.get("summary", ""))
        
        doc.add_heading("Skills", level=1)
        try:
            skills = json.loads(data.get("skills", "[]"))
            for skill in skills:
                doc.add_paragraph(f"{skill.get('category', '')}: {skill.get('items', '')}")
        except:
            doc.add_paragraph(data.get("skills", ""))
        
        doc.add_heading("Experience", level=1)
        try:
            experience = json.loads(data.get("experience", "[]"))
            for job in experience:
                doc.add_heading(f"{job.get('role', '')} — {job.get('company', '')}", level=2)
                doc.add_paragraph(f"{job.get('dates', '')} | {job.get('location', '')}")
                for bullet in job.get("bullets", []):
                    doc.add_paragraph(bullet, style="List Bullet")
        except:
            doc.add_paragraph(data.get("experience", ""))

        doc.add_heading("Education", level=1)
        try:
            education = json.loads(data.get("education", "[]"))
            for edu in education:
                doc.add_paragraph(f"{edu.get('title', '')} | {edu.get('institution', '')} — {edu.get('date', '')}")
        except:
            doc.add_paragraph(data.get("education", ""))

        doc.add_heading("Languages", level=1)
        doc.add_paragraph(data.get("languages", ""))
    
    else:
        # Fallback dump for dynamic fields
        doc.add_heading(TEMPLATES.get(payload.template, {}).get("name", "Document"), 0)
        for key, val in data.items():
            doc.add_heading(key.replace("_", " ").title(), level=2)
            doc.add_paragraph(str(val))

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return StreamingResponse(
        doc_io,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={payload.template}.docx"}
    )

@app.post("/import-json")
async def import_json(file: UploadFile = File(...)):
    """Receives a JSON upload, parses it, and sends fields back to UI."""
    content = await file.read()
    data = json.loads(content)
    return data